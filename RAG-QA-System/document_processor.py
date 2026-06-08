import os
from PyPDF2 import PdfReader
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import OllamaEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.schema import Document as LangChainDocument
import pickle

def load_pdf(file_path):
    """加载PDF文件并提取文本"""
    text = ""
    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            text += page.extract_text() or ""
    except Exception as e:
        print(f"加载PDF失败 {file_path}: {e}")
    return text

def load_docx(file_path):
    """加载DOCX文件并提取文本"""
    text = ""
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        print(f"加载DOCX失败 {file_path}: {e}")
    return text

def load_document(file_path):
    """根据文件扩展名加载不同类型的文档"""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        return load_pdf(file_path)
    elif ext == '.docx':
        return load_docx(file_path)
    elif ext == '.txt':
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    else:
        print(f"不支持的文件格式: {ext}")
        return ""

def load_documents_from_folder(folder_path):
    """批量加载文件夹中的所有文档"""
    documents = []
    filenames = []
    
    if not os.path.exists(folder_path):
        print(f"文件夹不存在: {folder_path}")
        return documents, filenames
    
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if os.path.isfile(file_path):
            text = load_document(file_path)
            if text.strip():
                documents.append(text)
                filenames.append(filename)
                print(f"已加载: {filename}")
    
    return documents, filenames

def split_documents(documents, chunk_size=1000, chunk_overlap=200):
    """将文档文本分块"""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", " ", ""]
    )
    
    all_chunks = []
    for doc in documents:
        chunks = text_splitter.split_text(doc)
        all_chunks.extend(chunks)
    
    print(f"文档分块完成，共生成 {len(all_chunks)} 个文本块")
    return all_chunks

def create_vector_db(chunks, persist_directory="./faiss_db"):
    """创建并保存向量数据库"""
    embeddings = OllamaEmbeddings(model="nomic-embed-text")
    
    langchain_docs = [LangChainDocument(page_content=chunk) for chunk in chunks]
    
    db = FAISS.from_documents(langchain_docs, embeddings)
    
    os.makedirs(persist_directory, exist_ok=True)
    db.save_local(persist_directory)
    print(f"向量数据库已保存到: {persist_directory}")
    return db

def load_vector_db(persist_directory="./faiss_db"):
    """加载已保存的向量数据库"""
    embeddings = OllamaEmbeddings(model="nomic-embed-text")
    
    db = FAISS.load_local(persist_directory, embeddings, allow_dangerous_deserialization=True)
    return db

def search_similar(db, query, k=3):
    """检索与查询最相关的k个文本块"""
    results = db.similarity_search_with_score(query, k=k)
    
    return {
        "documents": [doc.page_content for doc, score in results],
        "scores": [score for doc, score in results]
    }

if __name__ == "__main__":
    print("=" * 50)
    print("文档处理模块测试")
    print("=" * 50)
    
    os.makedirs("docs", exist_ok=True)
    
    test_content = """自然语言处理（NLP）是人工智能的一个重要分支，
它涉及计算机与人类语言之间的交互。NLP技术包括文本分类、情感分析、
命名实体识别、机器翻译等多个方向。

近年来，预训练语言模型如BERT、GPT等极大推动了NLP领域的发展。
这些模型通过在大规模文本数据上进行预训练，可以学习到丰富的语言表示。

深度学习在NLP中的应用包括：
1. 循环神经网络（RNN）
2. 长短时记忆网络（LSTM）
3. Transformer架构
4. 注意力机制

NLP的应用场景非常广泛，包括智能客服、机器翻译、文本摘要、
信息抽取等。随着技术的发展，NLP系统的性能不断提升，
越来越多的应用场景被开发出来。"""
    
    with open("docs/test_nlp.txt", "w", encoding="utf-8") as f:
        f.write(test_content)
    
    print("\n1. 加载文档...")
    docs, filenames = load_documents_from_folder("docs")
    print(f"加载了 {len(docs)} 个文档")
    
    print("\n2. 文本分块...")
    chunks = split_documents(docs)
    
    print("\n3. 创建向量数据库...")
    db = create_vector_db(chunks)
    
    print("\n4. 测试检索...")
    results = search_similar(db, "什么是自然语言处理？", k=3)
    print(f"找到 {len(results['documents'])} 个相关文档")
    for i, doc in enumerate(results['documents'], 1):
        print(f"\n结果 {i}:")
        print(doc[:200] + "..." if len(doc) > 200 else doc)
    
    print("\n" + "=" * 50)
    print("测试完成!")
    print("=" * 50)